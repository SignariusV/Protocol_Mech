import os
import re
from os import listdir
from win32com import client as wc
from docx import Document
import datetime
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE


def convert_docx(path_in, path_out):
    w = wc.Dispatch('Word.Application')
    doc = w.Documents.Open(path_in)
    doc.SaveAs(path_out + '.docx', 16)
    doc.Close()
    w.Quit()


def check_metr(name, protocol_path):
    path = 'X:\ИЦ Омега\Метрология\Формы 1,2,3,4,6_Паспорт ИЦ Омега'
    if os.path.isdir(path):
        for file in listdir(path):
            if file.endswith('doc') and not file.startswith('~') and name in file:
                try:
                    convert_docx(f'{path}\{file}', name)
                except:
                    print('Не удалось найти нужные файлы на сервере')
                break
    if name == 'СИ':
        table = 3
    elif name == 'ИО':
        table = 4
    protocol = Document(protocol_path)
    # style = protocol.styles.add_style('Warning', WD_STYLE_TYPE.PARAGRAPH)
    # style.font.name = 'Arial'
    # style.font.size = Pt(14)
    # style.font.italic = True
    # style.font.highlight_color = WD_COLOR_INDEX.YELLOW
    metr_form = Document(f"C:\\Users\\Vertaev\\Documents\\{name}.docx")
    for row in protocol.tables[table].rows[1:]:
        for form_row in metr_form.tables[1].rows:
            if row.cells[1].text in form_row.cells[2].text and row.cells[2].text in form_row.cells[4].text:
                row.cells[4].text = re.findall(r"\d{2}\.\d{2}\.\d{4}", form_row.cells[7].text)[-1]
                row.cells[4].paragraphs[0].style = 'Табл. по центру обычный'
                if datetime.datetime.strptime(row.cells[4].text, "%d.%m.%Y").date() <= datetime.date.today():
                    row.cells[4].paragraphs[0].style = 'Warning'
                break
    protocol.save(protocol_path)


name = 'СИ'
protocol_path = 'Doc.docx'
check_metr(name, protocol_path=protocol_path)
